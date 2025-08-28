import streamlit as st
import pdfplumber
from dashscope import Generation
from docx import Document
import io

# ========== 配置 ==========
API_KEY = " "  # 替换为你自己的 DashScope API Key

# ========== 提取 PDF 文本 ==========
def extract_pdf_text(uploaded_file, max_pages=5):
    """
    只处理可选文本 PDF
    """
    uploaded_bytes = uploaded_file.read()  # 读取文件 bytes
    text = ""
    with pdfplumber.open(io.BytesIO(uploaded_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            if i >= max_pages:
                break
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

# ========== 调用 DashScope ==========
def call_dashscope(text, lang="中文"):
    if not text.strip():
        raise ValueError("PDF 提取内容为空，请确认 PDF 可选文本可读")

    if lang == "中文":
        prompt = f"""
请根据以下论文内容，输出：
1. 用中文写一份200字摘要。
2. 提取5-10个关键英文术语，并给出对应中文解释。
3. 用英文写一份100字摘要。

论文内容如下：
{text}
"""
    else:
        prompt = f"""
Please summarize the following paper content:
1. Write a 200-word abstract in English.
2. Extract 5-10 key Chinese academic terms and provide English explanations.
3. Write a 100-word abstract in Chinese.

Paper content:
{text}
"""

    # 使用最新版 DashScope SDK 调用
    response = Generation.call(
        model="qwen-plus",
        api_key=API_KEY,
        prompt=prompt,
        max_output_tokens=1500
    )
    return response['output']['text']

# ========== 保存为 Word ==========
def save_to_word(content, filename="output.docx"):
    doc = Document()
    doc.add_heading("学术翻译 & 多语言摘要助手结果", level=1)
    for line in content.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
    doc.save(filename)
    return filename

# ========== Streamlit UI ==========
st.set_page_config(page_title="学术翻译 & 多语言摘要助手", layout="wide")
st.title("📚 学术翻译 & 多语言摘要助手")
st.write("上传论文 PDF，选择目标语言，生成摘要和术语表，并可下载 Word 文件。")

uploaded_file = st.file_uploader("上传论文 PDF 文件", type=["pdf"])
lang = st.radio("选择翻译方向", ["中文", "英文"], index=0)

if uploaded_file and st.button("开始处理"):
    with st.spinner("正在提取文本，请稍候..."):
        try:
            text = extract_pdf_text(uploaded_file, max_pages=5)
            st.write(f"📄 提取文本长度: {len(text)}")  # 查看提取长度

            if not text.strip():
                st.error("❌ PDF 提取到文本为空，请确认 PDF 可选文本可读")
            else:
                result = call_dashscope(text, lang)
                st.success("✅ 处理完成！")
                st.markdown("### 📄 结果输出")
                st.write(result)

                # 导出 Word
                word_file = save_to_word(result)
                with open(word_file, "rb") as f:
                    st.download_button("📥 下载 Word 文件", f, file_name=word_file)

        except Exception as e:
            st.error(f"❌ 处理失败: {e}")
